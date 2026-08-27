#!/usr/bin/env python3
"""Regression tests for knowledge indexing and page-copy checks."""

from __future__ import annotations

import json
import tempfile
import unittest
from pathlib import Path

from docx import Document

import content_checks
import knowledge_index


class KnowledgeIndexTests(unittest.TestCase):
    def test_extract_docx_includes_table_cells(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "table.docx"
            document = Document()
            document.add_paragraph("普通段落")
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "核心维度"
            table.cell(0, 1).text = "项目优势"
            document.save(path)

            chunks, warnings = knowledge_index.extract_docx(path, "table.docx")
            text = "\n".join(chunk["text"] for chunk in chunks)

            self.assertIn("普通段落", text)
            self.assertIn("核心维度", text)
            self.assertIn("项目优势", text)
            self.assertEqual(warnings, [])

    def test_project_rules_win_equal_score_ties(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            index_path = Path(directory) / "index.json"
            payload = {
                "version": 1,
                "sources": {
                    "02_资料库/方法.docx": {
                        "chunks": [{"source": "02_资料库/方法.docx", "page": None, "text": "课程内容"}],
                    },
                    "03_项目规范/口径.docx": {
                        "chunks": [{"source": "03_项目规范/口径.docx", "page": None, "text": "课程内容"}],
                    },
                },
            }
            index_path.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")

            results = knowledge_index.search_index(index_path, "课程内容", top_k=2)

            self.assertEqual(results[0]["source"], "03_项目规范/口径.docx")


class ContentCheckTests(unittest.TestCase):
    def test_quality_pass_ignores_exact_length_and_structure(self) -> None:
        reference = "# 原标题\n\n- 参考句一\n- 参考句二"
        candidate = "# 新标题\n\n这是一段原创、自然且适合排版的新文案。"

        result = content_checks.check_page(reference, candidate)

        self.assertTrue(result["quality_passes"])

    def test_visible_length_ignores_markdown_formatting(self) -> None:
        self.assertEqual(content_checks.visible_length("# **标题**"), 2)
        self.assertEqual(content_checks.visible_length("- `剧本`\n- 分镜"), 4)

    def test_number_details_exclude_list_markers_and_classify_numbers(self) -> None:
        text = "1. 剧本\n2. 分镜\n3个创作步骤\n课程共10周"

        details = content_checks.number_details(text)

        self.assertEqual(
            details,
            [
                {"value": "3", "kind": "structure"},
                {"value": "10", "kind": "fact"},
            ],
        )

    def test_allows_general_commercial_conversion_wording(self) -> None:
        self.assertEqual(content_checks.risk_matches("聚焦商业变现路径"), [])

    def test_blocks_income_and_guaranteed_results(self) -> None:
        matches = content_checks.risk_matches("保证学会，月入过万，预计20万曝光量，轻松回本，副业翻倍，零基础接单自由")
        categories = {item["category"] for item in matches}
        self.assertIn("guaranteed_result", categories)
        self.assertIn("income_or_price", categories)
        self.assertIn("commercial_metric", categories)
        matched_text = {item["text"] for item in matches}
        self.assertIn("轻松回本", matched_text)
        self.assertIn("副业翻倍", matched_text)
        self.assertIn("零基础接单自由", matched_text)

    def test_reports_markdown_page_structure(self) -> None:
        reference = "# 入门方法\n\n先理解流程。\n\n- 剧本\n- 分镜"
        candidate = "# 创作步骤\n\n先明确主题。\n\n- 画面\n- 剪辑"

        result = content_checks.check_page(reference, candidate)

        self.assertEqual(result["reference_structure"]["title_visible_length"], 4)
        self.assertEqual(result["candidate_structure"]["title_visible_length"], 4)
        self.assertEqual(result["reference_structure"]["paragraph_count"], 3)
        self.assertEqual(result["candidate_structure"]["list_item_count"], 2)
        self.assertTrue(result["structure_matches"])


class SkillConfigurationTests(unittest.TestCase):
    def test_default_knowledgebase_location_is_portable(self) -> None:
        skill = (Path(__file__).resolve().parents[1] / "SKILL.md").read_text(encoding="utf-8")
        self.assertNotIn("/" + "Users/", skill)
        self.assertIn("优先使用用户指定路径", skill)
        self.assertIn("未指定时使用当前工作区中的知识库", skill)

    def test_stages_are_strictly_sequential_and_downstream_blockers_do_not_skip_work(self) -> None:
        root = Path(__file__).parent.parent
        skill = (root / "SKILL.md").read_text(encoding="utf-8")

        positions = [skill.index(f"{number}. **") for number in range(1, 6)]
        self.assertEqual(positions, sorted(positions))
        self.assertIn("不得跳步、提前检查后续阶段或因后续阻塞省略前序阶段", skill)
        self.assertIn("后续阶段的阻塞只暂停该阶段", skill)

    def test_image_api_key_is_checked_only_immediately_before_generation(self) -> None:
        root = Path(__file__).parent.parent
        combined = "\n".join(
            (root / path).read_text(encoding="utf-8")
            for path in ("SKILL.md", "references/workflow.md")
        )

        self.assertIn("GLOBALAI_API_KEY", combined)
        self.assertIn("仅在生图前检查", combined)
        self.assertIn("密钥缺失不影响扫描、提取、检索、仿写和质检", combined)

    def test_workflow_auto_advances_without_confirmation_gates(self) -> None:
        root = Path(__file__).parent.parent
        combined = "\n".join(
            (root / path).read_text(encoding="utf-8")
            for path in ("SKILL.md", "references/workflow.md")
        )

        for obsolete in (
            "用户确认后继续",
            "等待用户确认",
            "三次确认不可跳过",
            "用户确认最终文案后",
        ):
            self.assertNotIn(obsolete, combined)
        self.assertIn("自动继续", combined)

    def test_reference_copy_extraction_is_versioned_and_preserves_history(self) -> None:
        root = Path(__file__).parent.parent
        combined = "\n".join(
            (root / path).read_text(encoding="utf-8")
            for path in ("SKILL.md", "references/workflow.md")
        )

        self.assertIn("参考图文案提取-YYYY-MM-DD", combined)
        self.assertIn("保留已有提取文件", combined)
        self.assertNotIn("写入 `工作文件/参考图文案提取.md`", combined)

    def test_copy_stage_focuses_on_originality_quality_and_fluency(self) -> None:
        root = Path(__file__).parent.parent
        combined = "\n".join(
            (root / path).read_text(encoding="utf-8")
            for path in ("SKILL.md", "references/workflow.md", "references/prompts.md")
        )

        self.assertIn("原创度", combined)
        self.assertIn("语言风格", combined)
        self.assertIn("通顺", combined)
        self.assertNotIn("有效字符数与参考页一致", combined)
        self.assertNotIn("标题有效字符数、段落数、列表项数和非空行数是否一致", combined)

    def test_index_lives_in_project_work_files_not_knowledgebase(self) -> None:
        workflow = (Path(__file__).parent.parent / "references" / "workflow.md").read_text(encoding="utf-8")
        self.assertIn("工作文件/知识库索引.json", workflow)
        self.assertNotIn("知识库根目录的 `.xiaohongshu-ai-comic-images/index.json`", workflow)

    def test_ai_common_knowledge_has_a_narrow_labeled_fallback(self) -> None:
        root = Path(__file__).parent.parent
        files = [root / "SKILL.md", root / "references" / "workflow.md", root / "references" / "prompts.md"]
        combined = "\n".join(path.read_text(encoding="utf-8") for path in files)

        self.assertIn("AI补充", combined)
        self.assertIn("通用方法、经验建议和解释性内容", combined)
        for forbidden in ("事实数字", "课程权益", "平台政策", "人物经历", "收益", "效果承诺"):
            self.assertIn(forbidden, combined)
        self.assertNotIn("不得用 AI 常识补写事实", combined)


if __name__ == "__main__":
    unittest.main(verbosity=2)
